"""
Сервис для анализа файлов: извлечение текста из PDF/DOC и описание изображений
"""
import io
import logging
from pathlib import Path
from typing import Optional, Dict, Any
import PyPDF2
from docx import Document
from PIL import Image

logger = logging.getLogger(__name__)


class FileAnalysisService:
    """Сервис для анализа загруженных файлов"""

    @staticmethod
    def extract_text_from_pdf(file_bytes: bytes) -> str:
        """Извлекает текст из PDF файла"""
        try:
            pdf_file = io.BytesIO(file_bytes)
            pdf_reader = PyPDF2.PdfReader(pdf_file)
            text_parts = []
            
            for page_num, page in enumerate(pdf_reader.pages):
                try:
                    text = page.extract_text()
                    if text.strip():
                        text_parts.append(f"--- Страница {page_num + 1} ---\n{text}")
                except Exception as e:
                    logger.warning(f"Ошибка извлечения текста со страницы {page_num + 1}: {e}")
                    continue
            
            extracted_text = "\n\n".join(text_parts)
            logger.info(f"✅ Извлечено {len(extracted_text)} символов из PDF ({len(pdf_reader.pages)} страниц)")
            return extracted_text
        except Exception as e:
            logger.error(f"❌ Ошибка извлечения текста из PDF: {e}")
            raise ValueError(f"Не удалось извлечь текст из PDF: {str(e)}")

    @staticmethod
    def extract_text_from_docx(file_bytes: bytes) -> str:
        """Извлекает текст из DOCX файла"""
        try:
            doc_file = io.BytesIO(file_bytes)
            doc = Document(doc_file)
            text_parts = []
            
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(paragraph.text)
            
            # Также извлекаем текст из таблиц
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join([cell.text.strip() for cell in row.cells if cell.text.strip()])
                    if row_text:
                        text_parts.append(row_text)
            
            extracted_text = "\n".join(text_parts)
            logger.info(f"✅ Извлечено {len(extracted_text)} символов из DOCX")
            return extracted_text
        except Exception as e:
            logger.error(f"❌ Ошибка извлечения текста из DOCX: {e}")
            raise ValueError(f"Не удалось извлечь текст из DOCX: {str(e)}")

    @staticmethod
    def extract_text_from_doc(file_bytes: bytes) -> str:
        """Извлекает текст из DOC файла (старый формат)"""
        # DOC файлы требуют специальных библиотек (python-docx не поддерживает старый формат)
        # Для простоты возвращаем сообщение об ошибке
        logger.warning("⚠️ Старый формат DOC не поддерживается. Используйте DOCX.")
        raise ValueError("Старый формат DOC не поддерживается. Пожалуйста, конвертируйте файл в DOCX или PDF.")

    @staticmethod
    def analyze_image(file_bytes: bytes, filename: str, llm_service, mime_type: str = "image/jpeg") -> Optional[str]:
        """Анализирует изображение через LLM с поддержкой vision"""
        try:
            # Проверяем, что это изображение
            image = Image.open(io.BytesIO(file_bytes))
            image_format = image.format
            
            logger.info(f"🖼️ Анализ изображения: {filename} ({image_format}, {image.size})")
            
            # Конвертируем изображение в base64 для отправки в LLM
            import base64
            image_buffer = io.BytesIO()
            
            # Определяем MIME тип на основе формата изображения
            format_to_mime = {
                'JPEG': 'image/jpeg',
                'JPG': 'image/jpeg',
                'PNG': 'image/png',
                'GIF': 'image/gif',
                'BMP': 'image/bmp',
                'WEBP': 'image/webp'
            }
            
            actual_mime_type = format_to_mime.get(image_format, mime_type)
            
            # Конвертируем в RGB если нужно (для JPEG)
            if image_format in ['PNG', 'JPEG', 'JPG']:
                if image.mode != 'RGB':
                    image = image.convert('RGB')
                image.save(image_buffer, format='JPEG', quality=85)
                actual_mime_type = 'image/jpeg'
            else:
                image.save(image_buffer, format=image_format)
            
            image_base64 = base64.b64encode(image_buffer.getvalue()).decode('utf-8')
            
            # Формируем промпт для анализа изображения
            prompt = """Проанализируй это изображение и опиши его содержимое подробно. 
            Если на изображении есть текст, извлеки его полностью.
            Если это график, диаграмма или таблица, опиши данные и значения.
            Если это документ или скриншот, опиши основное содержание.
            Если это фото, опиши что на нем изображено.
            Ответ должен быть информативным и структурированным."""
            
            # Используем vision API для анализа изображения
            try:
                # Проверяем, поддерживает ли LLMService анализ изображений
                if hasattr(llm_service, 'analyze_image'):
                    analysis = llm_service.analyze_image(image_base64, prompt, actual_mime_type)
                else:
                    # Fallback: используем обычный chat completion с описанием
                    logger.warning("⚠️ LLMService не поддерживает анализ изображений напрямую")
                    analysis = "Изображение загружено. Для анализа требуется поддержка vision API."
            except Exception as e:
                logger.error(f"❌ Ошибка анализа изображения через LLM: {e}")
                import traceback
                traceback.print_exc()
                analysis = f"Изображение загружено ({image_format}, {image.size[0]}x{image.size[1]}px). Ошибка анализа: {str(e)}"
            
            return analysis
        except Exception as e:
            logger.error(f"❌ Ошибка обработки изображения: {e}")
            import traceback
            traceback.print_exc()
            raise ValueError(f"Не удалось обработать изображение: {str(e)}")

    @staticmethod
    def analyze_file(file_bytes: bytes, filename: str, mime_type: str, llm_service=None) -> Dict[str, Any]:
        """
        Анализирует файл в зависимости от его типа
        
        Returns:
            dict с ключами:
            - extracted_text: извлеченный текст (для PDF/DOC)
            - analysis_result: результат анализа (для изображений)
            - file_type: тип файла
        """
        result = {
            "extracted_text": None,
            "analysis_result": None,
            "file_type": None
        }
        
        # Определяем тип файла
        file_ext = Path(filename).suffix.lower()
        
        try:
            if mime_type == 'application/pdf' or file_ext == '.pdf':
                result["file_type"] = "pdf"
                result["extracted_text"] = FileAnalysisService.extract_text_from_pdf(file_bytes)
                
            elif mime_type == 'application/vnd.openxmlformats-officedocument.wordprocessingml.document' or file_ext == '.docx':
                result["file_type"] = "docx"
                result["extracted_text"] = FileAnalysisService.extract_text_from_docx(file_bytes)
                
            elif mime_type == 'application/msword' or file_ext == '.doc':
                result["file_type"] = "doc"
                result["extracted_text"] = FileAnalysisService.extract_text_from_doc(file_bytes)
                
            elif mime_type and mime_type.startswith('image/'):
                result["file_type"] = "image"
                if llm_service:
                    result["analysis_result"] = FileAnalysisService.analyze_image(file_bytes, filename, llm_service, mime_type)
                else:
                    result["analysis_result"] = "Изображение загружено. Анализ недоступен (LLM сервис не настроен)."
            else:
                logger.warning(f"⚠️ Неподдерживаемый тип файла: {mime_type} ({file_ext})")
                result["file_type"] = "unknown"
                
        except Exception as e:
            logger.error(f"❌ Ошибка анализа файла {filename}: {e}")
            result["error"] = str(e)
        
        return result

